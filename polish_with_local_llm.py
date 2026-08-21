"""
Passo OPCIONAL: usa um LLM local (via Ollama) para revisar/organizar a
transcrição gerada pelo transcribe_diarize.py — corrige pontuação, remove
repetições de fala e (opcionalmente) gera um resumo por falante.

Esse passo é separado do transcribe_diarize.py de propósito: a transcrição
e a diarização rodam na CPU, mas o Ollama pode usar sua GPU AMD (Radeon XTX)
para acelerar o LLM, caso você configure o Ollama com suporte a GPU.

Pré-requisitos:
    1. Instalar o Ollama para Windows: https://ollama.com/download
    2. Baixar um modelo, ex:  ollama pull llama3.1:8b
    3. Deixar o Ollama rodando (ele já roda como serviço após instalado)

Uso:
    python polish_with_local_llm.py --input saida/video.docx --output saida/video_revisado.docx
    python polish_with_local_llm.py --input saida --output saida_revisada   (pasta inteira)
"""

import argparse
import sys
from pathlib import Path

import requests
from docx import Document

OLLAMA_URL = "http://localhost:11434/api/generate"

PROMPT_TEMPLATE = """Você é um assistente que revisa transcrições de áudio em português.
Receberá abaixo uma fala de UMA pessoa (já identificada). Sua tarefa:
- Corrigir pontuação e erros óbvios de transcrição
- Remover repetições e vícios de fala ("é...", "tipo assim", etc.) SEM mudar o sentido
- NÃO resumir, NÃO inventar conteúdo, NÃO traduzir
- Responder APENAS com o texto revisado, sem comentários

Texto original:
\"\"\"{text}\"\"\"

Texto revisado:"""


def polish_text(text: str, model: str) -> str:
    prompt = PROMPT_TEMPLATE.format(text=text)
    try:
        resp = requests.post(
            OLLAMA_URL,
            json={"model": model, "prompt": prompt, "stream": False},
            timeout=120,
        )
        resp.raise_for_status()
        return resp.json().get("response", text).strip()
    except requests.exceptions.ConnectionError:
        print("ERRO: não foi possível conectar ao Ollama em http://localhost:11434")
        print("Verifique se o Ollama está instalado e rodando.")
        sys.exit(1)


def polish_docx(input_path: Path, output_path: Path, model: str) -> None:
    doc = Document(str(input_path))
    new_doc = Document()

    for para in doc.paragraphs:
        if not para.runs:
            new_doc.add_paragraph(para.text)
            continue

        # O primeiro "run" é o rótulo do falante (ex: "[00:01:23] Pessoa 1: "), em negrito
        first_run = para.runs[0]
        if first_run.bold and ":" in first_run.text:
            speaker_label = first_run.text
            original_text = "".join(r.text for r in para.runs[1:])
            if original_text.strip():
                revised = polish_text(original_text, model)
            else:
                revised = original_text
            p = new_doc.add_paragraph()
            r1 = p.add_run(speaker_label)
            r1.bold = True
            p.add_run(revised)
        else:
            new_doc.add_paragraph(para.text)

    new_doc.save(str(output_path))
    print(f"Documento revisado salvo em: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Revisa transcrições .docx usando um LLM local (Ollama).")
    parser.add_argument("--input", required=True, help="Arquivo .docx ou pasta com vários .docx")
    parser.add_argument("--output", required=True, help="Arquivo .docx de saída ou pasta de saída")
    parser.add_argument("--model", default="llama3.1:8b", help="Nome do modelo no Ollama (padrão: llama3.1:8b)")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if input_path.is_dir():
        output_path.mkdir(parents=True, exist_ok=True)
        docx_files = sorted(input_path.glob("*.docx"))
        if not docx_files:
            print("Nenhum .docx encontrado em --input.")
            sys.exit(1)
        for f in docx_files:
            polish_docx(f, output_path / f.name, args.model)
    else:
        polish_docx(input_path, output_path, args.model)


if __name__ == "__main__":
    main()
