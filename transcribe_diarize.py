"""
Transcrição de vídeos .mkv com separação de falantes (diarização)
e geração automática de um documento Word (.docx) com o resultado.

Pipeline (100% local, CPU):
  1. Extrai o áudio do .mkv com ffmpeg
  2. Detecta "quem fala quando" com pyannote.audio (diarização)
  3. Transcreve o áudio com faster-whisper
  4. Cruza os dois resultados para rotular cada trecho com "Pessoa 1", "Pessoa 2"...
  5. Gera um .docx por vídeo

Uso básico:
    python transcribe_diarize.py --input video.mkv --output saida --hf-token SEU_TOKEN

Processar uma pasta inteira de vídeos:
    python transcribe_diarize.py --input pasta_com_videos --output saida --hf-token SEU_TOKEN

Dica: em vez de passar --hf-token toda vez, defina a variável de ambiente HF_TOKEN
(veja o README.md para como criar esse token).
"""

import argparse
import os
import subprocess
import sys
import tempfile
from pathlib import Path

from faster_whisper import WhisperModel
from pyannote.audio import Pipeline
import torch
from docx import Document


# --------------------------------------------------------------------------
# Áudio
# --------------------------------------------------------------------------

def extract_audio(video_path: Path, wav_path: Path) -> None:
    """Extrai o áudio do vídeo como WAV mono 16kHz (formato que os modelos esperam)."""
    cmd = [
        "ffmpeg", "-y", "-i", str(video_path),
        "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
        str(wav_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(
            f"ffmpeg falhou ao processar {video_path.name}:\n{result.stderr[-1500:]}"
        )


# --------------------------------------------------------------------------
# Diarização (separar quem fala)
# --------------------------------------------------------------------------

def run_diarization(pipeline: Pipeline, wav_path: Path):
    """Retorna lista de (inicio_seg, fim_seg, rotulo_falante)."""
    diarization = pipeline(str(wav_path))
    turns = []
    for turn, _, speaker in diarization.itertracks(yield_label=True):
        turns.append((turn.start, turn.end, speaker))
    return turns


def assign_speaker(seg_start: float, seg_end: float, turns) -> str:
    """Escolhe o falante cujo intervalo mais se sobrepõe ao trecho transcrito."""
    best_speaker, best_overlap = None, 0.0
    for t_start, t_end, spk in turns:
        overlap = min(seg_end, t_end) - max(seg_start, t_start)
        if overlap > best_overlap:
            best_overlap = overlap
            best_speaker = spk
    return best_speaker or "Desconhecido"


# --------------------------------------------------------------------------
# Transcrição
# --------------------------------------------------------------------------

def transcribe(model: WhisperModel, wav_path: Path, language: str | None):
    segments, info = model.transcribe(
        str(wav_path),
        language=language,          # None = detectar automaticamente
        vad_filter=True,            # ignora silêncios longos
        word_timestamps=False,
    )
    return list(segments), info


# --------------------------------------------------------------------------
# Combinar transcrição + diarização
# --------------------------------------------------------------------------

def merge_transcript_with_speakers(segments, turns):
    merged = []
    for seg in segments:
        speaker = assign_speaker(seg.start, seg.end, turns)
        text = seg.text.strip()
        if text:
            merged.append({"start": seg.start, "end": seg.end, "speaker": speaker, "text": text})
    return merged


def group_by_speaker_blocks(merged):
    """Junta trechos consecutivos do mesmo falante num único bloco de texto."""
    blocks = []
    for item in merged:
        if blocks and blocks[-1]["speaker"] == item["speaker"]:
            blocks[-1]["end"] = item["end"]
            blocks[-1]["text"] += " " + item["text"]
        else:
            blocks.append(dict(item))
    return blocks


def rename_speakers(blocks):
    """Troca rótulos internos (SPEAKER_00...) por 'Pessoa 1', 'Pessoa 2'... na ordem de aparição."""
    mapping = {}
    for b in blocks:
        if b["speaker"] not in mapping:
            mapping[b["speaker"]] = f"Pessoa {len(mapping) + 1}"
    for b in blocks:
        b["speaker"] = mapping[b["speaker"]]
    return blocks


def format_timestamp(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


# --------------------------------------------------------------------------
# Geração do .docx
# --------------------------------------------------------------------------

def build_docx(blocks, video_name: str, output_path: Path) -> None:
    doc = Document()
    doc.add_heading(f"Transcrição: {video_name}", level=1)

    if not blocks:
        doc.add_paragraph("(Nenhuma fala detectada.)")
    else:
        n_speakers = len({b["speaker"] for b in blocks})
        doc.add_paragraph(f"Falantes detectados: {n_speakers}")
        doc.add_paragraph("")

        for b in blocks:
            p = doc.add_paragraph()
            run_speaker = p.add_run(f"[{format_timestamp(b['start'])}] {b['speaker']}: ")
            run_speaker.bold = True
            p.add_run(b["text"])

    doc.save(str(output_path))


# --------------------------------------------------------------------------
# Processamento de um vídeo
# --------------------------------------------------------------------------

def process_video(video_path: Path, output_dir: Path, whisper_model, diar_pipeline, language):
    print(f"\n=== Processando: {video_path.name} ===")
    with tempfile.TemporaryDirectory() as tmp:
        wav_path = Path(tmp) / "audio.wav"

        print("  Extraindo áudio...")
        extract_audio(video_path, wav_path)

        print("  Separando falantes (diarização, pode demorar)...")
        turns = run_diarization(diar_pipeline, wav_path)
        n_falantes = len({s for _, _, s in turns})
        print(f"    -> {n_falantes} falante(s) detectado(s)")

        print("  Transcrevendo áudio...")
        segments, info = transcribe(whisper_model, wav_path, language)
        print(f"    -> idioma: {info.language} (confiança {info.language_probability:.2f})")

        merged = merge_transcript_with_speakers(segments, turns)
        blocks = group_by_speaker_blocks(merged)
        blocks = rename_speakers(blocks)

    output_path = output_dir / f"{video_path.stem}.docx"
    build_docx(blocks, video_path.name, output_path)
    print(f"  Documento salvo em: {output_path}")


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Transcreve vídeos .mkv separando falantes e gera um .docx por vídeo."
    )
    parser.add_argument("--input", required=True, help="Arquivo .mkv ou pasta com vários .mkv")
    parser.add_argument("--output", default="saida", help="Pasta onde os .docx serão salvos")
    parser.add_argument(
        "--model", default="medium",
        choices=["tiny", "base", "small", "medium", "large-v3"],
        help="Tamanho do modelo Whisper (maior = mais lento e mais preciso). Padrão: medium",
    )
    parser.add_argument(
        "--language", default="pt",
        help="Idioma do áudio (ex: pt, en). Use 'auto' para detectar automaticamente.",
    )
    parser.add_argument(
        "--hf-token", default=os.environ.get("HF_TOKEN"),
        help="Token do Hugging Face (necessário para baixar o modelo de diarização). "
             "Pode também ser definido na variável de ambiente HF_TOKEN.",
    )
    args = parser.parse_args()

    if not args.hf_token:
        print(
            "ERRO: é necessário um token do Hugging Face.\n"
            "Passe com --hf-token SEU_TOKEN ou defina a variável de ambiente HF_TOKEN.\n"
            "Veja o README.md para instruções de como criar o token."
        )
        sys.exit(1)

    language = None if args.language.lower() == "auto" else args.language

    input_path = Path(args.input)
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    if input_path.is_dir():
        videos = sorted(input_path.glob("*.mkv"))
    elif input_path.is_file():
        videos = [input_path]
    else:
        print(f"ERRO: caminho não encontrado: {input_path}")
        sys.exit(1)

    if not videos:
        print("Nenhum arquivo .mkv encontrado em --input.")
        sys.exit(1)

    print(f"{len(videos)} vídeo(s) para processar.")

    print("Carregando modelo Whisper (CPU, int8)...")
    whisper_model = WhisperModel(args.model, device="cpu", compute_type="int8")

    print("Carregando modelo de diarização (CPU)...")
    diar_pipeline = Pipeline.from_pretrained(
        "pyannote/speaker-diarization-3.1", use_auth_token=args.hf_token
    )
    if diar_pipeline is None:
        print(
            "\nERRO: não foi possível baixar o modelo de diarização.\n"
            "Confirme que você aceitou os termos de uso (com a MESMA conta do token) em:\n"
            "  https://huggingface.co/pyannote/speaker-diarization-3.1\n"
            "  https://huggingface.co/pyannote/segmentation-3.0\n"
            "e que o --hf-token informado é válido."
        )
        sys.exit(1)
    diar_pipeline.to(torch.device("cpu"))

    for video in videos:
        try:
            process_video(video, output_dir, whisper_model, diar_pipeline, language)
        except Exception as e:
            print(f"  ERRO ao processar {video.name}: {e}")

    print("\nConcluído.")


if __name__ == "__main__":
    main()